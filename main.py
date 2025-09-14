"""
Medical Device Compliance Scaffold - Optimized Version
Integrates LangGraph workflows, Autogen auditors, and document generation
for ISO 13485, ISO 14971, and IEC 62304 compliance

Token usage optimized and PEP-compliant version
"""

import asyncio
import json
import os
from collections import defaultdict
from dataclasses import dataclass, field
from datetime import datetime
from enum import Enum
from pathlib import Path
from typing import Any, Dict, List, Optional, TypedDict, Union

import autogen
import pandas as pd
import PyPDF2
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from dotenv import load_dotenv
from langgraph.graph import END, START, StateGraph
from langgraph.graph.message import add_messages

# ============================================================================
# CONSTANTS AND CONFIGURATION
# ============================================================================

# Token usage optimization constants
MAX_TOKENS_PER_REQUEST = 4000
MAX_CONTEXT_LENGTH = 8000
DEFAULT_TEMPERATURE = 0.1

# File processing constants
SUPPORTED_EXTENSIONS = {'.docx', '.pdf', '.txt'}
MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB
MIN_DOCUMENT_LENGTH = 500

# Output file templates
DOCX_TEMPLATE = "{standard}_compliance_document.docx"
EXCEL_TEMPLATE = "traceability_matrix_{timestamp}.xlsx"


# ============================================================================
# CORE DATA MODELS
# ============================================================================

class StandardType(Enum):
    """Medical device compliance standards."""
    ISO_13485 = "ISO 13485"
    ISO_14971 = "ISO 14971"  
    IEC_62304 = "IEC 62304"


class ComplianceStatus(Enum):
    """Compliance item status enumeration."""
    OPEN = "Open"
    IN_PROGRESS = "In Progress"
    COMPLETED = "Completed"
    BLOCKED = "Blocked"


class RiskLevel(Enum):
    """Risk level enumeration."""
    LOW = "Low"
    MEDIUM = "Medium"
    HIGH = "High"
    CRITICAL = "Critical"


@dataclass
class TraceabilityItem:
    """Represents a single item in the traceability matrix."""
    
    id: str
    standard: StandardType
    requirement: str
    implementation: str
    verification: str
    validation: str
    status: ComplianceStatus = ComplianceStatus.OPEN
    risk_level: RiskLevel = RiskLevel.MEDIUM
    assigned_to: str = ""
    completion_date: str = ""

    def to_dict(self) -> Dict[str, str]:
        """Convert to dictionary for export."""
        return {
            'ID': self.id,
            'Standard': self.standard.value,
            'Requirement': self.requirement,
            'Implementation': self.implementation,
            'Verification': self.verification,
            'Validation': self.validation,
            'Status': self.status.value,
            'Risk Level': self.risk_level.value,
            'Assigned To': self.assigned_to,
            'Completion Date': self.completion_date
        }


@dataclass
class ComplianceDocument:
    """Represents a compliance document."""
    
    title: str
    standard: StandardType
    content: str
    traceability_items: List[TraceabilityItem] = field(default_factory=list)
    audit_results: Dict[str, Any] = field(default_factory=dict)
    created_date: str = field(default_factory=lambda: datetime.now().isoformat())
    
    @property
    def word_count(self) -> int:
        """Get approximate word count of content."""
        return len(self.content.split())
    
    @property
    def is_valid_length(self) -> bool:
        """Check if document meets minimum length requirements."""
        return len(self.content) >= MIN_DOCUMENT_LENGTH


class WorkflowState(TypedDict):
    """LangGraph workflow state with comprehensive typing."""
    
    messages: List[Dict[str, Any]]
    current_document: Optional[ComplianceDocument]
    audit_feedback: Dict[str, List[str]]
    traceability_matrix: List[TraceabilityItem]
    workflow_step: str
    input_document_content: str
    token_usage: Dict[str, int]


# ============================================================================
# TOKEN USAGE OPTIMIZATION
# ============================================================================

class TokenOptimizer:
    """Handles token usage optimization for LLM calls."""
    
    def __init__(self):
        self.token_usage = defaultdict(int)
        self.context_cache = {}
    
    def estimate_tokens(self, text: str) -> int:
        """Estimate token count (rough approximation: 1 token ≈ 4 chars)."""
        return len(text) // 4
    
    def optimize_prompt(self, prompt: str, max_tokens: int = MAX_TOKENS_PER_REQUEST) -> str:
        """Optimize prompt to fit within token limits."""
        estimated_tokens = self.estimate_tokens(prompt)
        
        if estimated_tokens <= max_tokens:
            return prompt
        
        # Truncate from middle to preserve context structure
        words = prompt.split()
        target_words = (max_tokens * 4) // 5  # Conservative estimate
        
        if len(words) <= target_words:
            return prompt
        
        # Keep beginning and end, truncate middle
        keep_start = target_words // 3
        keep_end = target_words // 3
        
        truncated = (
            ' '.join(words[:keep_start]) +
            f'\n\n[... content truncated for token efficiency ...]\n\n' +
            ' '.join(words[-keep_end:])
        )
        
        return truncated
    
    def get_cached_response(self, prompt_hash: str) -> Optional[str]:
        """Retrieve cached response if available."""
        return self.context_cache.get(prompt_hash)
    
    def cache_response(self, prompt_hash: str, response: str) -> None:
        """Cache response for future use."""
        if len(self.context_cache) > 100:  # Simple cache size limit
            # Remove oldest entry
            oldest_key = next(iter(self.context_cache))
            del self.context_cache[oldest_key]
        
        self.context_cache[prompt_hash] = response


# ============================================================================
# DOCUMENT PROCESSING - PEP 8 COMPLIANT
# ============================================================================

class DocumentProcessor:
    """Handles document reading and processing with proper error handling."""
    
    @staticmethod
    def validate_file(file_path: Union[str, Path]) -> bool:
        """Validate file exists, size, and extension."""
        path = Path(file_path)
        
        if not path.exists():
            return False
        
        if path.stat().st_size > MAX_FILE_SIZE:
            return False
        
        if path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            return False
        
        return True
    
    @staticmethod
    def read_docx_file(file_path: Union[str, Path]) -> str:
        """Extract text content from DOCX file.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text content
            
        Raises:
            IOError: If file cannot be read
        """
        try:
            doc = Document(file_path)
            content_lines = []
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    content_lines.append(text)
            
            return '\n'.join(content_lines)
        
        except Exception as e:
            raise IOError(f"Error reading DOCX file {file_path}: {e}")
    
    @staticmethod
    def read_pdf_file(file_path: Union[str, Path]) -> str:
        """Extract text content from PDF file.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text content
            
        Raises:
            IOError: If file cannot be read
        """
        try:
            content_lines = []
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page in pdf_reader.pages:
                    text = page.extract_text().strip()
                    if text:
                        content_lines.append(text)
            
            return '\n'.join(content_lines)
        
        except Exception as e:
            raise IOError(f"Error reading PDF file {file_path}: {e}")
    
    @staticmethod 
    def read_text_file(file_path: Union[str, Path]) -> str:
        """Read plain text file.
        
        Args:
            file_path: Path to text file
            
        Returns:
            File content as string
            
        Raises:
            IOError: If file cannot be read
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            raise IOError(f"Error reading text file {file_path}: {e}")
    
    @classmethod
    def load_document(cls, file_path: Union[str, Path]) -> str:
        """Load document content based on file extension.
        
        Args:
            file_path: Path to document file
            
        Returns:
            Document content as string
            
        Raises:
            ValueError: If file format unsupported
            IOError: If file cannot be read
        """
        path = Path(file_path)
        
        if not cls.validate_file(path):
            raise ValueError(f"Invalid file: {file_path}")
        
        extension = path.suffix.lower()
        
        if extension == '.docx':
            return cls.read_docx_file(path)
        elif extension == '.pdf':
            return cls.read_pdf_file(path)
        elif extension == '.txt':
            return cls.read_text_file(path)
        else:
            raise ValueError(f"Unsupported file format: {extension}")


# ============================================================================
# OPTIMIZED AUDITOR SYSTEM
# ============================================================================

class OptimizedAuditorSystem:
    """Token-efficient auditor system with caching and batching."""
    
    def __init__(self, api_key: str):
        self.api_key = api_key
        self.token_optimizer = TokenOptimizer()
        self._create_auditors()
    
    def _create_auditors(self) -> None:
        """Create optimized auditor agents with concise system messages."""
        llm_config = {
            "config_list": [{
                "model": "gpt-4o-mini",  # Fixed model name
                "api_key": self.api_key
            }],
            "temperature": DEFAULT_TEMPERATURE,
        }
        
        # Optimized, concise system messages
        self.auditors = {
            "iso13485": autogen.AssistantAgent(
                name="ISO13485_Auditor",
                system_message=(
                    "ISO 13485 QMS auditor. Review for: QMS processes, "
                    "document control, management responsibility, "
                    "product realization. Provide specific clause references."
                ),
                llm_config=llm_config
            ),
            
            "iso14971": autogen.AssistantAgent(
                name="ISO14971_Auditor", 
                system_message=(
                    "ISO 14971 Risk Management auditor. Review for: risk process, "
                    "hazard analysis, risk controls, management file. "
                    "Focus on completeness and traceability."
                ),
                llm_config=llm_config
            ),
            
            "iec62304": autogen.AssistantAgent(
                name="IEC62304_Auditor",
                system_message=(
                    "IEC 62304 Software auditor. Review for: safety classification, "
                    "development planning, requirements, architecture, testing. "
                    "Ensure lifecycle compliance."
                ),
                llm_config=llm_config
            )
        }
    
    async def audit_document_batch(self, 
                                 document: ComplianceDocument,
                                 standards: List[StandardType]) -> Dict[str, List[str]]:
        """Audit document efficiently using batched requests."""
        # For demonstration, using simulated responses to avoid actual API calls
        # In production, this would make optimized API calls
        
        audit_feedback = {}
        
        optimized_content = self.token_optimizer.optimize_prompt(
            document.content, 
            MAX_TOKENS_PER_REQUEST
        )
        
        # Simulate efficient auditing process
        for standard in standards:
            standard_key = standard.value.lower().replace(' ', '').replace('-', '')
            audit_feedback[standard_key] = self._get_optimized_feedback(
                standard, 
                optimized_content
            )
        
        return audit_feedback
    
    def _get_optimized_feedback(self, 
                              standard: StandardType, 
                              content: str) -> List[str]:
        """Generate optimized feedback for a standard."""
        # This would be replaced with actual LLM calls in production
        feedback_map = {
            StandardType.ISO_13485: [
                "QMS process mapping requires more detail in section 4.1",
                "Document control procedure needs version management clarity",
                "Management review inputs should include customer feedback data"
            ],
            StandardType.ISO_14971: [
                "Risk analysis completeness: verify all use scenarios covered", 
                "Risk control measures need effectiveness validation criteria",
                "Post-market surveillance integration with risk management required"
            ],
            StandardType.IEC_62304: [
                "Software safety classification justification needs enhancement",
                "Architecture design documentation should include security considerations",
                "Verification activities must trace to all software requirements"
            ]
        }
        
        return feedback_map.get(standard, ["Standard-specific feedback needed"])


# ============================================================================
# IMPROVED DOCUMENT TEMPLATES
# ============================================================================

class ComplianceTemplates:
    """Optimized compliance document templates."""
    
    # Template constants for consistency
    COMMON_SECTIONS = [
        "Document Control",
        "Purpose and Scope", 
        "Definitions and Abbreviations",
        "Roles and Responsibilities"
    ]
    
    @staticmethod
    def get_iso13485_template() -> str:
        """Generate ISO 13485 compliant template."""
        return """# Quality Management System Procedure

## Document Control
Document ID: QMS-001
Version: 1.0
Effective Date: {date}
Next Review: {next_review}

## Purpose and Scope
Establish Quality Management System processes per ISO 13485:2016 
requirements for medical device development and manufacturing.

## Definitions and Abbreviations
- QMS: Quality Management System
- CAPA: Corrective and Preventive Action
- DHF: Design History File

## Roles and Responsibilities
- Management Representative: QMS oversight and regulatory liaison
- Quality Manager: Daily QMS operations and compliance monitoring
- Process Owners: Individual process control and improvement

## QMS Process Framework
The QMS encompasses interconnected processes from design control
through post-market surveillance, ensuring regulatory compliance
and continuous improvement.

## Process Controls and Documentation
All QMS processes are controlled through:
- Documented procedures and work instructions
- Process monitoring and measurement
- Regular management review and improvement

## Monitoring and Measurement
- Internal audits per planned schedule
- Management reviews quarterly
- Process performance indicators tracked monthly
- Customer satisfaction monitoring

## Improvement and CAPA
Systematic approach to:
- Identify improvement opportunities
- Implement corrective actions
- Prevent recurrence of nonconformities
- Verify effectiveness of actions
""".format(
            date=datetime.now().strftime("%Y-%m-%d"),
            next_review=(datetime.now().replace(year=datetime.now().year + 1)
                        .strftime("%Y-%m-%d"))
        )
    
    @staticmethod
    def get_iso14971_template() -> str:
        """Generate ISO 14971 compliant template."""
        return """# Risk Management Plan

## Document Control  
Document ID: RMP-001
Version: 1.0
Effective Date: {date}
Device: [Insert Device Name]

## Purpose and Scope
Establish risk management process per ISO 14971:2019 for [Device Name]
throughout product lifecycle from concept to disposal.

## Risk Management Policy
Commitment to systematic risk management ensuring:
- Patient and user safety
- Risk-benefit analysis
- Regulatory compliance
- Continuous risk monitoring

## Risk Management Process
1. Risk Analysis: Identify hazards and hazardous situations  
2. Risk Evaluation: Assess risk acceptability
3. Risk Control: Implement control measures
4. Risk Management File: Document all activities
5. Post-Market Surveillance: Monitor residual risks

## Risk Acceptability Criteria
- High Risk: Requires multiple independent controls
- Medium Risk: Requires control measures with verification
- Low Risk: May be acceptable as-is with documentation

## Risk Control Hierarchy
1. Inherent safety by design
2. Protective measures and alarms
3. Information for safety (labeling, training)

## Risk Management File Contents
- Risk management plan
- Hazard analysis and risk assessment  
- Risk control implementation and verification
- Risk-benefit analysis
- Post-market surveillance plan and reports

## Review and Updates
Plan reviewed annually and updated for:
- Design changes
- New hazards identified
- Post-market feedback
- Regulatory changes
""".format(date=datetime.now().strftime("%Y-%m-%d"))
    
    @staticmethod
    def get_iec62304_template() -> str:
        """Generate IEC 62304 compliant template."""
        return """# Software Development Plan

## Document Control
Document ID: SDP-001
Version: 1.0  
Effective Date: {date}
Software Item: [Insert Software Name]

## Software Description and Intended Use
[Software overview, intended medical purpose, user environment,
and integration with medical device system]

## Safety Classification
Software Safety Classification per IEC 62304:
- Class A: Non-life-threatening
- Class B: Non-life-threatening injury possible  
- Class C: Death or serious injury possible

Classification: [Insert Class with Justification]

## Development Lifecycle Model
Lifecycle model: [V-Model/Agile/Waterfall]
Justification: [Rationale for model selection]

## Development Process Overview  
1. Planning (Section 5.1)
2. Requirements Analysis (Section 5.2)
3. Architectural Design (Section 5.3)  
4. Detailed Design (Section 5.4)
5. Implementation (Section 5.5)
6. Integration and Testing (Section 5.6)
7. System Testing (Section 5.7)
8. Release (Section 5.8)

## Requirements Management
- Requirements traceability matrix maintained
- Requirements verification methods defined
- Change control process established

## Verification and Validation Strategy
- Unit testing for all software units
- Integration testing for software interfaces
- System testing against requirements
- Validation in intended use environment

## Risk Control Measures
Software implements risk control measures identified in
risk management process (ISO 14971), including:
- Input validation and error handling
- Alarm and alert systems
- Data integrity verification

## Configuration Management
- Version control system for all software items
- Build and release procedures
- Change control and approval process

## Problem Resolution Process
- Bug tracking and resolution
- Impact analysis for changes
- Regression testing requirements
""".format(date=datetime.now().strftime("%Y-%m-%d"))
    
    @classmethod
    def get_template(cls, standard: StandardType) -> str:
        """Get template for specified standard."""
        template_map = {
            StandardType.ISO_13485: cls.get_iso13485_template,
            StandardType.ISO_14971: cls.get_iso14971_template,
            StandardType.IEC_62304: cls.get_iec62304_template
        }
        
        template_func = template_map.get(standard)
        if template_func:
            return template_func()
        
        return "# Compliance Document Template\n\nTemplate not available."


# ============================================================================
# SAMPLE DATA GENERATION - OPTIMIZED
# ============================================================================

def create_comprehensive_traceability_matrix() -> List[TraceabilityItem]:
    """Create comprehensive traceability matrix with proper status distribution."""
    
    matrix_data = [
        # ISO 13485 Items
        {
            "id": "ISO13485-4.1.1",
            "standard": StandardType.ISO_13485,
            "requirement": "QMS process establishment and documentation",
            "implementation": "Process map QM-001 with interaction matrix",
            "verification": "QA review and management approval completed",
            "validation": "QMS audit results demonstrate effectiveness",
            "status": ComplianceStatus.COMPLETED,
            "risk_level": RiskLevel.HIGH
        },
        {
            "id": "ISO13485-4.2.3", 
            "standard": StandardType.ISO_13485,
            "requirement": "Document control for current versions",
            "implementation": "Electronic document system DCS-v2.1 deployed",
            "verification": "System testing and user acceptance completed",
            "validation": "Document control effectiveness verified in audit",
            "status": ComplianceStatus.COMPLETED,
            "risk_level": RiskLevel.MEDIUM
        },
        {
            "id": "ISO13485-7.3.2",
            "standard": StandardType.ISO_13485, 
            "requirement": "Design and development planning",
            "implementation": "Design control procedure DC-001 in development",
            "verification": "Procedure draft under review",
            "validation": "Validation pending procedure approval",
            "status": ComplianceStatus.IN_PROGRESS,
            "risk_level": RiskLevel.HIGH
        },
        
        # ISO 14971 Items  
        {
            "id": "ISO14971-4.1",
            "standard": StandardType.ISO_14971,
            "requirement": "Risk management process establishment", 
            "implementation": "Risk management plan RMP-001 approved",
            "verification": "Process documentation reviewed by risk committee",
            "validation": "Process effectiveness demonstrated through application",
            "status": ComplianceStatus.COMPLETED,
            "risk_level": RiskLevel.CRITICAL
        },
        {
            "id": "ISO14971-4.4", 
            "standard": StandardType.ISO_14971,
            "requirement": "Comprehensive risk analysis execution",
            "implementation": "FMEA completed for all subsystems and interfaces",
            "verification": "Risk committee review and approval completed",
            "validation": "Risk analysis validation in progress",
            "status": ComplianceStatus.IN_PROGRESS,
            "risk_level": RiskLevel.HIGH
        },
        {
            "id": "ISO14971-5.2",
            "standard": StandardType.ISO_14971,
            "requirement": "Risk control measure implementation",
            "implementation": "Control measures design phase",
            "verification": "Design review scheduled",
            "validation": "Validation planning in progress", 
            "status": ComplianceStatus.OPEN,
            "risk_level": RiskLevel.HIGH
        },
        
        # IEC 62304 Items
        {
            "id": "IEC62304-4.3",
            "standard": StandardType.IEC_62304,
            "requirement": "Software safety classification",
            "implementation": "Class B classification documented with rationale",
            "verification": "Classification review completed by safety team",
            "validation": "Classification validated against risk analysis",
            "status": ComplianceStatus.COMPLETED,
            "risk_level": RiskLevel.CRITICAL
        },
        {
            "id": "IEC62304-5.1.1",
            "standard": StandardType.IEC_62304,
            "requirement": "Software development process planning",
            "implementation": "Software development plan SDP-001 in review",
            "verification": "Plan technical review scheduled",
            "validation": "Process validation pending plan approval",
            "status": ComplianceStatus.IN_PROGRESS,
            "risk_level": RiskLevel.MEDIUM  
        },
        {
            "id": "IEC62304-5.5.1",
            "standard": StandardType.IEC_62304,
            "requirement": "Software integration and testing",
            "implementation": "Integration test strategy drafted",
            "verification": "Strategy under technical review",
            "validation": "Integration validation planning required",
            "status": ComplianceStatus.OPEN,
            "risk_level": RiskLevel.MEDIUM
        }
    ]
    
    return [TraceabilityItem(**item) for item in matrix_data]


# ============================================================================
# OPTIMIZED WORKFLOW IMPLEMENTATION  
# ============================================================================

class OptimizedComplianceWorkflow:
    """Token-optimized compliance workflow with enhanced error handling."""
    
    def __init__(self):
        self.token_optimizer = TokenOptimizer()
        self.document_processor = DocumentProcessor()
        self.templates = ComplianceTemplates()
    
    def create_workflow(self) -> StateGraph:
        """Create optimized LangGraph workflow."""
        
        def initialize_workflow(state: WorkflowState) -> WorkflowState:
            """Initialize workflow with comprehensive setup."""
            print("🚀 Initializing optimized compliance workflow...")
            
            state["traceability_matrix"] = create_comprehensive_traceability_matrix()
            state["audit_feedback"] = {}
            state["workflow_step"] = "initialized"
            state["token_usage"] = defaultdict(int)
            
            state["messages"] = add_messages(state["messages"], [{
                "role": "system",
                "content": (
                    f"Workflow initialized with {len(state['traceability_matrix'])} "
                    "traceability items across all standards"
                )
            }])
            
            return state
        
        def generate_document(state: WorkflowState) -> WorkflowState:
            """Generate optimized compliance document."""
            print("📄 Generating compliance document...")
            
            # Determine standard - default to ISO 13485 if not specified
            standard = StandardType.ISO_13485
            
            # Use input content or template
            if state.get("input_document_content"):
                content = state["input_document_content"]
                print("📥 Using uploaded document content")
                
                # Optimize content for token efficiency
                content = self.token_optimizer.optimize_prompt(
                    content, 
                    MAX_CONTEXT_LENGTH
                )
            else:
                content = self.templates.get_template(standard)
                print("📋 Using optimized template")
            
            # Filter relevant traceability items
            relevant_items = [
                item for item in state["traceability_matrix"]
                if item.standard == standard
            ]
            
            # Create document with validation
            doc = ComplianceDocument(
                title=f"{standard.value} Compliance Document",
                standard=standard,
                content=content,
                traceability_items=relevant_items
            )
            
            # Validate document
            if not doc.is_valid_length:
                print(f"⚠️ Document below minimum length ({len(doc.content)} chars)")
            
            state["current_document"] = doc
            state["workflow_step"] = "document_generated"
            state["token_usage"]["document_generation"] = self.token_optimizer.estimate_tokens(content)
            
            state["messages"] = add_messages(state["messages"], [{
                "role": "assistant", 
                "content": (
                    f"Generated {standard.value} document: "
                    f"{doc.word_count} words, {len(relevant_items)} traceability items"
                )
            }])
            
            return state
        
        def audit_document(state: WorkflowState) -> WorkflowState:
            """Perform optimized document audit."""
            print("🔍 Auditing document with optimized feedback...")
            
            # Simulate optimized audit with realistic feedback
            doc = state["current_document"]
            
            # Token-efficient audit feedback
            audit_feedback = {
                "iso13485": [
                    f"QMS process clarity: {doc.word_count} words may need structure review",
                    "Management review frequency specification recommended",
                    "CAPA integration with process controls needs enhancement"
                ],
                "iso14971": [
                    "Risk evaluation criteria require quantitative thresholds",
                    "Post-market surveillance integration plan needed",
                    "Risk control effectiveness measurement methods required"
                ],
                "iec62304": [
                    "Software architecture documentation depth assessment needed",
                    "Integration testing coverage verification required", 
                    "Requirements traceability matrix completeness check needed"
                ],
                "cross_standard": [
                    "Risk management integration with QMS processes",
                    "Software verification alignment with risk controls",
                    "Traceability consistency across all standards"
                ]
            }
            
            state["audit_feedback"] = audit_feedback
            state["workflow_step"] = "document_audited"
            
            # Calculate token usage for audit
            feedback_text = " ".join([
                " ".join(feedback) 
                for feedback in audit_feedback.values()
            ])
            state["token_usage"]["audit_feedback"] = self.token_optimizer.estimate_tokens(feedback_text)
            
            total_issues = sum(len(feedback) for feedback in audit_feedback.values())
            
            state["messages"] = add_messages(state["messages"], [{
                "role": "system",
                "content": f"Audit completed: {total_issues} recommendations across all standards"
            }])
            
            return state
        
        def finalize_compliance(state: WorkflowState) -> WorkflowState:
            """Finalize compliance documentation with optimizations."""
            print("✅ Finalizing compliance documentation...")
            
            doc = state["current_document"]
            if not doc:
                raise ValueError("No document available for finalization")
            
            doc.audit_results = state["audit_feedback"]
            
            # Generate safe filename
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            safe_standard = doc.standard.value.replace(' ', '_').replace('-', '')
            output_path = f"{safe_standard}_compliance_{timestamp}.docx"
            
            try:
                self._generate_optimized_docx(doc, output_path)
                print(f"📝 Document generated: {output_path}")
            except Exception as e:
                print(f"❌ Document generation error: {e}")
                output_path = "document_generation_failed.txt"
            
            state["workflow_step"] = "completed"
            
            # Calculate total token usage
            total_tokens = sum(state["token_usage"].values())
            
            state["messages"] = add_messages(state["messages"], [{
                "role": "assistant",
                "content": (
                    f"Compliance workflow completed. "
                    f"Output: {output_path}, Token usage: {total_tokens}"
                )
            }])
            
            return state
        
        # Build optimized workflow graph
        workflow = StateGraph(WorkflowState)
        
        workflow.add_node("initialize", initialize_workflow)
        workflow.add_node("generate_document", generate_document) 
        workflow.add_node("audit_document", audit_document)
        workflow.add_node("finalize", finalize_compliance)
        
        workflow.add_edge(START, "initialize")
        workflow.add_edge("initialize", "generate_document")
        workflow.add_edge("generate_document", "audit_document")
        workflow.add_edge("audit_document", "finalize")
        workflow.add_edge("finalize", END)
        
        return workflow.compile()
    
    def _generate_optimized_docx(self, doc: ComplianceDocument, output_path: str) -> None:
        """Generate Word document with enhanced formatting and validation."""
        try:
            word_doc = Document()
            
            # Title with proper formatting
            title = word_doc.add_heading(doc.title, 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Document metadata table
            self._add_metadata_table(word_doc, doc)
            word_doc.add_paragraph()
            
            # Compliance validation results
            issues = self._validate_document_compliance(doc)
            if issues:
                self._add_compliance_issues(word_doc, issues)
                word_doc.add_page_break()
            
            # Main content with proper formatting
            self._add_formatted_content(word_doc, doc.content)
            
            # Traceability matrix
            if doc.traceability_items:
                word_doc.add_page_break()
                self._add_traceability_matrix(word_doc, doc.traceability_items)
            
            # Audit results summary
            if doc.audit_results:
                word_doc.add_page_break() 
                self._add_audit_results(word_doc, doc.audit_results)
            
            # Save with error handling
            word_doc.save(output_path)
            
        except Exception as e:
            raise IOError(f"Failed to generate document {output_path}: {e}")
    
    def _add_metadata_table(self, doc: Document, compliance_doc: ComplianceDocument) -> None:
        """Add formatted metadata table."""
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Table Grid'
        
        metadata = [
            ("Standard", compliance_doc.standard.value),
            ("Created Date", compliance_doc.created_date[:10]),
            ("Word Count", str(compliance_doc.word_count)),
            ("Traceability Items", str(len(compliance_doc.traceability_items))),
            ("Document Status", "Under Review" if compliance_doc.audit_results else "Draft")
        ]
        
        for i, (label, value) in enumerate(metadata):
            table.cell(i, 0).text = label
            table.cell(i, 1).text = value
    
    def _validate_document_compliance(self, doc: ComplianceDocument) -> List[str]:
        """Validate document compliance with enhanced checks."""
        issues = []
        
        if not doc.is_valid_length:
            issues.append(f"Document too brief: {len(doc.content)} chars (min: {MIN_DOCUMENT_LENGTH})")
        
        if "TODO" in doc.content or "TBD" in doc.content:
            issues.append("Document contains placeholder content")
        
        if "[Insert" in doc.content:
            issues.append("Document contains template placeholders")
        
        required_sections = {
            StandardType.ISO_13485: ["Purpose", "QMS", "Process", "Control"],
            StandardType.ISO_14971: ["Risk", "Analysis", "Control", "Management"],
            StandardType.IEC_62304: ["Software", "Development", "Safety", "Testing"]
        }
        
        content_lower = doc.content.lower()
        missing_concepts = []
        
        for concept in required_sections.get(doc.standard, []):
            if concept.lower() not in content_lower:
                missing_concepts.append(concept)
        
        if missing_concepts:
            issues.append(f"Missing key concepts: {', '.join(missing_concepts)}")
        
        return issues
    
    def _add_compliance_issues(self, doc: Document, issues: List[str]) -> None:
        """Add compliance issues section."""
        doc.add_heading('Compliance Review Items', level=1)
        
        for issue in issues:
            p = doc.add_paragraph()
            p.style = 'List Bullet'
            p.text = issue
    
    def _add_formatted_content(self, doc: Document, content: str) -> None:
        """Add main content with proper formatting."""
        doc.add_heading('Document Content', level=1)
        
        lines = content.strip().split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            else:
                doc.add_paragraph(line)
    
    def _add_traceability_matrix(self, doc: Document, items: List[TraceabilityItem]) -> None:
        """Add formatted traceability matrix."""
        doc.add_heading('Traceability Matrix', level=1)
        
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        
        headers = ['ID', 'Requirement', 'Implementation', 'Status', 'Risk', 'Verification']
        for i, header in enumerate(headers):
            table.cell(0, i).text = header
        
        for item in items:
            row_cells = table.add_row().cells
            row_cells[0].text = item.id
            row_cells[1].text = self._truncate_text(item.requirement, 50)
            row_cells[2].text = self._truncate_text(item.implementation, 50) 
            row_cells[3].text = item.status.value
            row_cells[4].text = item.risk_level.value
            row_cells[5].text = self._truncate_text(item.verification, 50)
    
    def _add_audit_results(self, doc: Document, audit_results: Dict[str, List[str]]) -> None:
        """Add audit results summary."""
        doc.add_heading('Audit Results Summary', level=1)
        
        for auditor, feedback_list in audit_results.items():
            doc.add_heading(auditor.replace('_', ' ').title(), level=2)
            
            for feedback in feedback_list:
                p = doc.add_paragraph()
                p.style = 'List Bullet'
                p.text = feedback
    
    @staticmethod
    def _truncate_text(text: str, max_length: int) -> str:
        """Truncate text with ellipsis if too long."""
        return f"{text[:max_length]}..." if len(text) > max_length else text


# ============================================================================
# ENHANCED REPORTING AND EXPORT
# ============================================================================

class ComplianceReporter:
    """Enhanced compliance reporting with token optimization."""
    
    @staticmethod
    def export_traceability_matrix(items: List[TraceabilityItem], 
                                 output_path: Optional[str] = None) -> str:
        """Export traceability matrix to Excel with enhanced formatting."""
        if not output_path:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = f"traceability_matrix_{timestamp}.xlsx"
        
        # Convert to DataFrame
        data = [item.to_dict() for item in items]
        df = pd.DataFrame(data)
        
        with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name='Traceability Matrix', index=False)
            
            # Format columns
            worksheet = writer.sheets['Traceability Matrix']
            
            for column in worksheet.columns:
                max_length = 0
                column_letter = column[0].column_letter
                
                for cell in column:
                    try:
                        cell_length = len(str(cell.value))
                        if cell_length > max_length:
                            max_length = cell_length
                    except (AttributeError, TypeError):
                        pass
                
                # Set column width with reasonable limits
                adjusted_width = min(max_length + 2, 50)
                worksheet.column_dimensions[column_letter].width = adjusted_width
        
        print(f"📊 Traceability matrix exported: {output_path}")
        return output_path
    
    @staticmethod
    def generate_compliance_report(items: List[TraceabilityItem]) -> Dict[str, Any]:
        """Generate comprehensive compliance status report."""
        if not items:
            return {"error": "No traceability items provided"}
        
        report = {
            "summary": {
                "total_items": len(items),
                "generated_date": datetime.now().isoformat(),
            },
            "by_standard": {},
            "by_status": defaultdict(int),
            "by_risk_level": defaultdict(int),
            "completion_metrics": {}
        }
        
        # Calculate statistics
        completed_count = 0
        high_risk_count = 0
        
        for item in items:
            # By standard
            standard_name = item.standard.value
            if standard_name not in report["by_standard"]:
                report["by_standard"][standard_name] = {
                    "total": 0,
                    "completed": 0,
                    "in_progress": 0,
                    "open": 0
                }
            
            report["by_standard"][standard_name]["total"] += 1
            
            # Track status
            if item.status == ComplianceStatus.COMPLETED:
                completed_count += 1
                report["by_standard"][standard_name]["completed"] += 1
            elif item.status == ComplianceStatus.IN_PROGRESS:
                report["by_standard"][standard_name]["in_progress"] += 1
            else:
                report["by_standard"][standard_name]["open"] += 1
            
            # Overall status and risk tracking
            report["by_status"][item.status.value] += 1
            report["by_risk_level"][item.risk_level.value] += 1
            
            if item.risk_level in [RiskLevel.HIGH, RiskLevel.CRITICAL]:
                high_risk_count += 1
        
        # Convert defaultdicts to regular dicts
        report["by_status"] = dict(report["by_status"])
        report["by_risk_level"] = dict(report["by_risk_level"])
        
        # Calculate completion metrics
        total_items = len(items)
        completion_percentage = (completed_count / total_items) * 100
        
        report["completion_metrics"] = {
            "completion_percentage": round(completion_percentage, 1),
            "completed_count": completed_count,
            "remaining_count": total_items - completed_count,
            "high_risk_count": high_risk_count,
            "high_risk_percentage": round((high_risk_count / total_items) * 100, 1)
        }
        
        return report
    
    @staticmethod
    def print_report_summary(report: Dict[str, Any]) -> None:
        """Print formatted compliance report summary."""
        print("\n" + "=" * 80)
        print("COMPLIANCE REPORT SUMMARY")
        print("=" * 80)
        
        summary = report["summary"]
        metrics = report["completion_metrics"]
        
        print(f"📊 Total Items: {summary['total_items']}")
        print(f"✅ Completion: {metrics['completion_percentage']}%")
        print(f"🎯 Completed: {metrics['completed_count']}")
        print(f"⏳ Remaining: {metrics['remaining_count']}")
        print(f"⚠️  High Risk: {metrics['high_risk_count']} ({metrics['high_risk_percentage']}%)")
        
        print(f"\n📋 BY STANDARD:")
        for standard, stats in report["by_standard"].items():
            completion = (stats['completed'] / stats['total']) * 100
            print(f"   {standard}: {stats['total']} items ({completion:.1f}% complete)")
        
        print(f"\n🎯 BY STATUS:")
        for status, count in report["by_status"].items():
            print(f"   {status}: {count} items")
        
        print(f"\n⚠️  BY RISK LEVEL:")
        for risk_level, count in report["by_risk_level"].items():
            print(f"   {risk_level}: {count} items")


# ============================================================================
# MAIN EXECUTION FUNCTION - OPTIMIZED
# ============================================================================

async def run_optimized_compliance_workflow(input_file_path: Optional[str] = None) -> WorkflowState:
    """Execute optimized compliance workflow with comprehensive error handling."""
    
    print("=" * 80)
    print("OPTIMIZED MEDICAL DEVICE COMPLIANCE WORKFLOW")
    print("Token Usage Optimized | PEP 8 Compliant")  
    print("=" * 80)
    
    # Initialize workflow system
    workflow_system = OptimizedComplianceWorkflow()
    reporter = ComplianceReporter()
    
    # Load input document if provided
    input_content = ""
    if input_file_path:
        try:
            print(f"📂 Loading document: {input_file_path}")
            input_content = DocumentProcessor.load_document(input_file_path)
            print(f"✅ Document loaded: {len(input_content):,} characters")
        except (ValueError, IOError) as e:
            print(f"⚠️ Document loading failed: {e}")
            print("📋 Continuing with template-based workflow")
    
    # Execute workflow
    workflow = workflow_system.create_workflow()
    
    initial_state: WorkflowState = {
        "messages": [],
        "current_document": None,
        "audit_feedback": {},
        "traceability_matrix": [],
        "workflow_step": "starting",
        "input_document_content": input_content,
        "token_usage": {}
    }
    
    try:
        final_state = await workflow.ainvoke(initial_state)
        
        print("\n" + "=" * 80)
        print("WORKFLOW EXECUTION RESULTS")
        print("=" * 80)
        
        # Display workflow messages
        for msg in final_state["messages"]:
            content = getattr(msg, 'content', str(msg))
            print(f"📝 {content}")
        
        # Export traceability matrix
        matrix_file = reporter.export_traceability_matrix(final_state["traceability_matrix"])
        
        # Generate and display compliance report
        report = reporter.generate_compliance_report(final_state["traceability_matrix"])
        reporter.print_report_summary(report)
        
        # Display token usage summary
        token_usage = final_state.get("token_usage", {})
        if token_usage:
            total_tokens = sum(token_usage.values())
            print(f"\n🔧 TOKEN USAGE SUMMARY:")
            print(f"   Total Estimated Tokens: {total_tokens:,}")
            for operation, tokens in token_usage.items():
                print(f"   {operation.replace('_', ' ').title()}: {tokens:,} tokens")
        
        print(f"\n✅ GENERATED FILES:")
        if final_state.get('current_document'):
            doc_title = final_state['current_document'].title.replace(' ', '_')
            print(f"   - {doc_title}.docx (compliance document)")
        print(f"   - {matrix_file} (traceability matrix)")
        
        return final_state
        
    except Exception as e:
        print(f"❌ Workflow execution failed: {e}")
        raise


# ============================================================================
# MAIN ENTRY POINT
# ============================================================================

def main():
    """Main entry point with argument handling."""
    import sys
    
    load_dotenv()
    
    # Simple argument handling
    input_file = None
    if len(sys.argv) > 1:
        input_file = sys.argv[1]
        if not Path(input_file).exists():
            print(f"❌ File not found: {input_file}")
            sys.exit(1)
    
    try:
        # Run the optimized workflow
        asyncio.run(run_optimized_compliance_workflow(input_file))
    except KeyboardInterrupt:
        print("\n⏹️ Workflow interrupted by user")
    except Exception as e:
        print(f"❌ Workflow failed: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()
